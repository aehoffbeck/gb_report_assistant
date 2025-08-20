from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import List, Optional
import os
from docx import Document
from datetime import datetime

app = FastAPI()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# --------------------------
# Endpoint 1: Get template fields and prompts
# --------------------------
class TemplateRequest(BaseModel):
    clientName: str
    templateName: str

@app.post("/report-template-fields")
def get_template_fields(payload: TemplateRequest):
    templates = {
        "competitor analysis": {
            "fieldNames": [
                "clientName",
                "dateSending",
                "analyzedDateRange",
                "reportObjective",
                "numberOfCompetitors",
                "competitorInclusionReason",
                "competitorsAnalyzed"
            ],
            "userPrompts": [
                "What is the client's name?",
                "What date will this report be sent?",
                "What date range is being analyzed?",
                "What are you looking to gain from this report?",
                "How many competitors are being analyzed?",
                "Why did you include these specific competitors?",
                "List the competitors being analyzed."
            ]
        }
    }

    template_key = payload.templateName.strip().lower()
    template_info = templates.get(template_key)

    if not template_info:
        return JSONResponse(status_code=400, content={
            "error": f"Unsupported template '{payload.templateName}'."
        })

    uploads = [
        {
            "type": "graph",
            "source": "Meltwater",
            "description": "Upload exported graphs from Meltwater related to this report."
        },
        {
            "type": "spreadsheet",
            "format": ["csv", "xlsx"],
            "description": "Upload the media coverage spreadsheet for this period."
        }
    ]

    return {
        "clientName": payload.clientName,
        "template": payload.templateName,
        "fieldNames": template_info["fieldNames"],
        "userPrompts": template_info["userPrompts"],
        "requiresUploads": uploads
    }

# --------------------------
# Placeholder: Analyze graphs
# --------------------------
@app.post("/analyze-graphs")
async def analyze_graphs(files: List[UploadFile] = File(...)):
    return {
        "insights": [
            "Spike in mentions during launch.",
            "Positive sentiment overall.",
            "Top engagement from Instagram."
        ]
    }

# --------------------------
# Placeholder: Analyze spreadsheet
# --------------------------
@app.post("/analyze-spreadsheet")
async def analyze_spreadsheet(file: UploadFile = File(...)):
    return {
        "insights": [
            "TechCrunch and Wired were top sources.",
            "Coverage spike on July 10.",
            "Sentiment 65% positive overall."
        ]
    }

# --------------------------
# Generate executive summary
# --------------------------
class ExecSummaryRequest(BaseModel):
    clientName: str
    templateName: str
    fieldResponses: dict
    graphInsights: Optional[List[str]] = []
    spreadsheetInsights: Optional[List[str]] = []

@app.post("/generate-executive-summary")
def generate_summary(data: ExecSummaryRequest):
    bullets = []
    for k, v in data.fieldResponses.items():
        bullets.append(f"- {k}: {v}")
    bullets += [f"• {i}" for i in (data.graphInsights + data.spreadsheetInsights)]

    summary = f"Executive Summary for {data.clientName}:\n\n" + "\n".join(bullets)

    return {
        "clientName": data.clientName,
        "executiveSummary": summary
    }

# --------------------------
# Generate final .docx report
# --------------------------
@app.post("/generate-report-docx")
async def generate_report_docx(
    clientName: str = Form(...),
    templateName: str = Form(...),
    executiveSummary: str = Form(...),
    keyFindings: List[str] = Form(...),
    graphFiles: List[UploadFile] = File(default=[])
):
    template_file = f"templates/{templateName.replace(' ', '_').lower()}_template.docx"
    if not os.path.exists(template_file):
        return JSONResponse(status_code=404, content={"error": "Template not found"})

    doc = Document(template_file)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executiveSummary)

    doc.add_heading("Key Findings", level=1)
    for finding in keyFindings:
        doc.add_paragraph(f"- {finding}", style="List Bullet")

    output_path = f"{UPLOAD_DIR}/{clientName.replace(' ', '_')}_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(output_path)

    return FileResponse(output_path, filename=os.path.basename(output_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
